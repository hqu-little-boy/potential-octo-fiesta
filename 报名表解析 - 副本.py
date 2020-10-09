import csv
from docx import Document
import os
import shutil

def SignUpDocx(TITLE,i,path):#文件生成
    document = Document()
    document.add_heading("报名表", level=0)
    for count in range(1,12):
        document.add_heading(TITLE[count], level=1,)
        document.add_paragraph(i[count],style="Intense Quote")
    document.save(PATH+path+i[1]+".docx")#可以覆盖保存，PATH是根目录

def CreateFolder(path,name_list):#文件夹生成
    for i in name_list:#['第一志愿/','第二志愿/']:
        isExists = os.path.exists(path+str(i))
        if not isExists:
            os.makedirs(path+str(i))	
            print("{} 目录创建成功".format(i))
        else:
            print("{} 目录已经存在，文件删除出错，请检查".format(i))
            break

def DuplicateRemoval(prime_list):#根据学号去重
    for i in prime_list:
        for j in prime_list[:prime_list.index(i)]:
            if j[3] == i[3]:
                prime_list.remove(j)
    return prime_list

def Shunt(wish,TITLE,sign_up_list):#分流
    for i in sign_up_list[1:]:
        SignUpDocx(TITLE,i,wish+i[TITLE.index(wish[:-1])]+"/")#wish[:-1]是去掉/

PATH = "F:/报名表/"
f = open(PATH + "学生社团联合会志愿者电子报名表.csv", "r", encoding="UTF-8")
csvreader = csv.reader(f)
sign_up_list = DuplicateRemoval(list(csvreader))#去重
TITLE = sign_up_list[0]

shutil.rmtree(PATH+'第一志愿/')
shutil.rmtree(PATH+'第二志愿/')#删除非空文件夹，上同
CreateFolder(PATH,['第一志愿/','第二志愿/'])#删除之前的数据，已去掉重复填报者

for wish in ['第一志愿/','第二志愿/']:
    CreateFolder(PATH + wish,["宣传媒体中心","办公室","社团活动管理部","就业部","财务与监察部","社团事务部","阳光服务部","自律会"])
    Shunt(wish,TITLE,sign_up_list)
